import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from datetime import date

# Fungsi untuk memaksa font menjadi Times New Roman & Warna Hitam
def set_font_times(run, size=11, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = None # Memastikan warna hitam (bukan biru default)
    
    # Kompatibilitas font untuk MS Word
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = qn('w:eastAsia')
    rPr.xpath('./w:rFonts')[0].set(rFonts, 'Times New Roman')
    rPr.xpath('./w:rFonts')[0].set(qn('w:ascii'), 'Times New Roman')
    rPr.xpath('./w:rFonts')[0].set(qn('w:hAnsi'), 'Times New Roman')

def generate_word(proyek, tanggal, tempat, team, st1, st2):
    doc = Document()
    
    # --- HEADER SECTION ---
    # Daily Report (Paling Atas)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_times(p1.add_run('DAILY REPORT'), size=14, bold=True)
    
    # Monitoring IP Network (Tengah)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_times(p2.add_run('MONITORING IP NETWORK'), size=12, bold=True)
    
    # Universitas Jember (Bawah)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_times(p3.add_run('UNIVERSITAS JEMBER'), size=12, bold=True)
    
    # Garis Pembatas Sederhana
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_times(p_line.add_run("-" * 50))

    # --- INFORMASI PROYEK (Rata Kiri, Titik Dua Nempel) ---
    items = [
        (f"Tanggal: {tanggal.strftime('%d %B %Y')}"),
        (f"Tempat: {tempat}"),
        (f"Project: {proyek}"),
        (f"Team On Site: {team}")
    ]

    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font_times(run)
        # Mengatur spasi antar baris agar tidak terlalu jauh
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph() # Jeda sedikit

    # --- CHECKLIST MONITORING ---
    p_check = doc.add_paragraph()
    run_check = p_check.add_run('CHECKLIST MONITORING')
    set_font_times(run_check, bold=True)

    chk_table = doc.add_table(rows=3, cols=2)
    chk_table.style = 'Table Grid'
    
    # Header Tabel
    cells_h = chk_table.rows[0].cells
    set_font_times(cells_h[0].paragraphs[0].add_run("Perangkat / IP"), bold=True)
    set_font_times(cells_h[1].paragraphs[0].add_run("Status"), bold=True)
    
    # Isi Tabel
    row1 = chk_table.rows[1].cells
    set_font_times(row1[0].paragraphs[0].add_run("PTP Linknet (139.225.2.201)"))
    set_font_times(row1[1].paragraphs[0].add_run(st1))
    
    row2 = chk_table.rows[2].cells
    set_font_times(row2[0].paragraphs[0].add_run("PTP Unej (139.255.2.202)"))
    set_font_times(row2[1].paragraphs[0].add_run(st2))

    # Simpan ke memori
    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- INTERFACE STREAMLIT ---
st.set_page_config(page_title="Network Report UNEJ", layout="centered")

st.markdown("### 📋 Daily Report Generator")

with st.form("main_form"):
    proyek = st.text_input("Project", value="Monitoring IP Network UNEJ")
    tanggal = st.date_input("Pilih Tanggal", date.today())
    tempat = st.text_input("Tempat", value="Universitas Jember")
    team = st.text_input("Team On Site", value="Anggi")

    st.markdown("---")
    st.write("**Status Konektivitas**")
    c1, c2 = st.columns(2)
    with c1:
        st1 = st.radio("PTP Linknet", ["Up", "Down"], horizontal=True)
    with c2:
        st2 = st.radio("PTP Unej", ["Up", "Down"], horizontal=True)
    
    submit = st.form_submit_button("Generate & Download Word")

if submit:
    word_file = generate_word(proyek, tanggal, tempat, team, st1, st2)
    st.download_button(
        label="💾 Download Laporan",
        data=word_file,
        file_name=f"Report_{tanggal}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )